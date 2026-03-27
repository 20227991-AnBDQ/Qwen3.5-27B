import os
import json
import requests
import configparser
from pathlib import Path
import PyPDF2
from docx import Document

# ============================================
# PHẦN 1: XỬ LÝ FILE PDF
# ============================================

def extract_text_from_pdf(pdf_path):
    """Trích text từ file PDF"""
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        print(f"  ❌ Lỗi PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_path):
    """Trích text từ file DOCX (bao gồm bảng, hình ảnh text)"""
    try:
        doc = Document(docx_path)
        text = ""
        
        # Trích text từ paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"
        
        # Trích text từ bảng (nếu có)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # Trích text từ shapes (text boxes)
        try:
            for shape in doc.element.body.iter():
                if hasattr(shape, 'text') and shape.text.strip():
                    text += shape.text + "\n"
        except:
            pass
        
        return text.strip()
    except Exception as e:
        print(f"  ❌ Lỗi DOCX: {str(e)}")
        return ""

# ============================================
# PHẦN 2: XỬ LÝ FILE .URL
# ============================================

def extract_url_from_url_file(url_file_path):
    """Trích URL từ file .url (Windows shortcut)"""
    try:
        config = configparser.ConfigParser()
        config.read(url_file_path)
        
        if 'InternetShortcut' in config:
            url = config['InternetShortcut'].get('URL')
            return url
    except Exception as e:
        print(f"  ❌ Lỗi trích URL: {str(e)}")
    
    return None

def download_cv_from_url(url, output_path, timeout=30):
    """Tải CV từ URL"""
    try:
        # Thêm headers để giả lập browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, timeout=timeout, headers=headers, allow_redirects=True)
        
        if response.status_code == 200:
            # Kiểm tra nếu là PDF
            if 'application/pdf' in response.headers.get('content-type', ''):
                with open(output_path, 'wb') as f:
                    f.write(response.content)
                return True
            # Kiểm tra nếu là HTML (TopCV, etc.)
            elif 'text/html' in response.headers.get('content-type', ''):
                print(f"  ⚠️  URL là trang web (HTML), không phải PDF trực tiếp")
                print(f"     URL: {url[:60]}...")
                return False
            else:
                with open(output_path, 'wb') as f:
                    f.write(response.content)
                return True
        else:
            print(f"  ❌ HTTP {response.status_code}")
            return False
    except requests.exceptions.Timeout:
        print(f"  ❌ Timeout (URL quá chậm)")
        return False
    except requests.exceptions.ConnectionError:
        print(f"  ❌ Lỗi kết nối (URL không tồn tại hoặc bị chặn)")
        return False
    except Exception as e:
        print(f"  ❌ Lỗi: {str(e)[:50]}")
        return False

# ============================================
# PHẦN 3: PIPELINE CHÍNH
# ============================================

def process_mixed_cvs(input_dir="data/raw_cvs", output_file="data/cvs_extracted.jsonl"):
    """
    Xử lý folder chứa cả PDF và .url files
    
    Quy trình:
    1. Quét folder tìm .url files
    2. Trích URL từ .url files
    3. Tải CV từ URL
    4. Trích text từ PDF/DOCX
    5. Lưu thành JSONL
    """
    
    os.makedirs(input_dir, exist_ok=True)
    
    cvs_data = []
    pdf_count = 0
    url_count = 0
    failed_count = 0
    
    print("="*60)
    print("🔄 XỬ LÝ CV (PDF + .URL)")
    print("="*60)
    
    # ========== BƯỚC 1: XỬ LÝ FILE .URL ==========
    print("\n📥 BƯỚC 1: Xử lý file .url...")
    print("-"*60)
    
    url_files = list(Path(input_dir).glob("*.url"))
    
    if url_files:
        print(f"Tìm thấy {len(url_files)} file .url\n")
        
        for url_file in url_files:
            print(f"  📄 {url_file.name}")
            
            # Trích URL
            url = extract_url_from_url_file(str(url_file))
            
            if not url:
                print(f"    ❌ Không trích được URL")
                failed_count += 1
                continue
            
            print(f"    🔗 URL: {url[:50]}...")
            
            # Tải CV
            cv_filename = f"downloaded_{url_count}_{url_file.stem}.pdf"
            cv_path = os.path.join(input_dir, cv_filename)
            
            print(f"    ⬇️  Tải xuống...")
            
            if download_cv_from_url(url, cv_path):
                print(f"    ✅ Tải thành công: {cv_filename}")
                url_count += 1
            else:
                print(f"    ❌ Tải thất bại")
                failed_count += 1
    else:
        print("  ℹ️  Không tìm thấy file .url")
    
    # ========== BƯỚC 2: XỬ LÝ FILE PDF/DOCX ==========
    print("\n📄 BƯỚC 2: Trích text từ PDF/DOCX...")
    print("-"*60)
    
    # Tìm tất cả file PDF và DOCX (bao gồm trong subfolder)
    pdf_files = list(Path(input_dir).rglob("*.pdf"))
    docx_files = list(Path(input_dir).rglob("*.docx"))
    
    all_files = pdf_files + docx_files
    
    print(f"Tìm thấy {len(pdf_files)} PDF + {len(docx_files)} DOCX\n")
    
    for i, file_path in enumerate(all_files, 1):
        print(f"  [{i}/{len(all_files)}] {file_path.name}")
        
        # Trích text
        if file_path.suffix.lower() == '.pdf':
            text = extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            text = extract_text_from_docx(str(file_path))
        else:
            text = ""
        
        if text.strip() and len(text) > 50:  # Kiểm tra text có ý nghĩa
            cvs_data.append({
                "filename": file_path.name,
                "text": text,
                "source": "pdf" if file_path.suffix.lower() == '.pdf' else "docx",
                "file_path": str(file_path)
            })
            print(f"    ✅ Trích thành công ({len(text)} ký tự)")
            pdf_count += 1
        else:
            print(f"    ❌ Không trích được text hoặc text quá ngắn")
            failed_count += 1
    
    # ========== BƯỚC 3: LƯU JSONL ==========
    print("\n💾 BƯỚC 3: Lưu dữ liệu...")
    print("-"*60)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        for cv in cvs_data:
            f.write(json.dumps(cv, ensure_ascii=False) + "\n")
    
    # ========== THỐNG KÊ ==========
    print("\n" + "="*60)
    print("✅ HOÀN THÀNH")
    print("="*60)
    print(f"📊 Thống kê:")
    print(f"  - PDF/DOCX trích thành công: {pdf_count}")
    print(f"  - .URL tải thành công: {url_count}")
    print(f"  - Thất bại: {failed_count}")
    print(f"  - Tổng cộng: {len(cvs_data)} CV")
    print(f"\n💾 Lưu vào: {output_file}")
    
    # Thống kê chi tiết
    pdf_success = sum(1 for cv in cvs_data if cv['source'] == 'pdf')
    docx_success = sum(1 for cv in cvs_data if cv['source'] == 'docx')
    
    print(f"\n📈 Chi tiết:")
    print(f"  - PDF: {pdf_success}")
    print(f"  - DOCX: {docx_success}")
    print(f"  - .URL: {url_count}")
    
    return cvs_data

# ============================================
# CHẠY SCRIPT
# ============================================

if __name__ == "__main__":
    # Cài đặt thư viện nếu cần
    print("📦 Kiểm tra thư viện...")
    try:
        import PyPDF2
        from docx import Document
        print("✅ Thư viện sẵn sàng\n")
    except ImportError:
        print("⚠️  Cài đặt thư viện:")
        print("  pip install PyPDF2 python-docx requests\n")
    
    # Chạy pipeline
    cvs = process_mixed_cvs(
        input_dir="data/raw_cvs",
        output_file="data/cvs_extracted.jsonl"
    )
    
    print("\n" + "="*60)
    print("🎯 Bước tiếp theo:")
    print("  python normalize_cv.py")
    print("="*60)
